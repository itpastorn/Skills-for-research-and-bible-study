#!/usr/bin/env python3
"""
Build a .docx report from a markdown draft, with IBM Plex Sans applied
throughout.

Usage:
    python build_docx.py --input draft.md --output report.docx

Dependencies:
    - pandoc (system binary)
    - python-docx
    - lxml

The script runs pandoc to produce a raw .docx, then walks the document
and sets IBM Plex Sans on every run, every style's rPr, and on the
document defaults. This is how Lars wants Word documents produced by
this skill.
"""

import argparse
import subprocess
import sys
import tempfile
from pathlib import Path

TARGET_FONT = "IBM Plex Sans"
NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def run_pandoc(md_path: Path, out_path: Path) -> None:
    """Convert markdown to docx via pandoc."""
    cmd = [
        "pandoc",
        str(md_path),
        "-o",
        str(out_path),
        "--from",
        "markdown",
        "--to",
        "docx",
    ]
    subprocess.run(cmd, check=True)


def apply_font(docx_path: Path, out_path: Path) -> None:
    """Apply TARGET_FONT to all runs, styles, and docDefaults."""
    from docx import Document
    from lxml import etree

    doc = Document(str(docx_path))

    def set_rfonts(rpr):
        # Drop any existing rFonts then add a fresh one with the target.
        for rfonts in rpr.findall(f"{NS}rFonts"):
            rpr.remove(rfonts)
        rfonts = etree.SubElement(rpr, f"{NS}rFonts")
        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
            rfonts.set(f"{NS}{attr}", TARGET_FONT)

    # 1. All styles
    for style in doc.styles:
        element = style.element
        rpr = element.find(f"{NS}rPr")
        if rpr is None:
            rpr = etree.SubElement(element, f"{NS}rPr")
        set_rfonts(rpr)

    # 2. All runs in all paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = TARGET_FONT

    # 3. Runs inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = TARGET_FONT

    # 4. docDefaults
    styles_root = doc.styles.element
    doc_defaults = styles_root.find(f"{NS}docDefaults")
    if doc_defaults is not None:
        rpr_default = doc_defaults.find(f"{NS}rPrDefault")
        if rpr_default is not None:
            rpr = rpr_default.find(f"{NS}rPr")
            if rpr is None:
                rpr = etree.SubElement(rpr_default, f"{NS}rPr")
        set_rfonts(rpr)

    doc.save(str(out_path))


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, help="Path to markdown draft")
    parser.add_argument("--output", required=True, help="Path to final .docx")
    args = parser.parse_args()

    md_path = Path(args.input)
    out_path = Path(args.output)

    if not md_path.exists():
        print(f"error: input not found: {md_path}", file=sys.stderr)
        return 1

    out_path.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory() as tmpdir:
        raw_path = Path(tmpdir) / "raw.docx"
        run_pandoc(md_path, raw_path)
        apply_font(raw_path, out_path)

    print(f'wrote {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
